"""
md_to_docx.py — converts GDD_Julian_Gomez.md → GDD_Julian_Gomez.docx
Auto-installs python-docx if missing. Run: python md_to_docx.py
"""
import subprocess, sys, re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx …")
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Paths — try several candidates so the script works regardless of cwd
# ---------------------------------------------------------------------------
OUT_FILE = Path(__file__).parent / "GDD_Julian_Gomez.docx"

_candidates = [
    # Hardcoded absolute path (known from project Glob)
    Path(r"C:\Users\Teilnehmer\Desktop\Schule\PRG\Abschlussprojekt_SRH_26\Dokumente_Schule\Completed\GDD_Julian_Gomez.md"),
    # Relative: script is in docs/Dokumente_Schule/Completed/
    Path(__file__).parent.parent.parent.parent / "Dokumente_Schule" / "Completed" / "GDD_Julian_Gomez.md",
    Path(__file__).parent.parent.parent / "Dokumente_Schule" / "Completed" / "GDD_Julian_Gomez.md",
    # Same folder as the docx (docs/Dokumente_Schule/Completed/)
    Path(__file__).parent / "GDD_Julian_Gomez.md",
]

print("Searching for GDD markdown …")
MD_FILE = None
for c in _candidates:
    exists = c.exists()
    print(f"  {'FOUND' if exists else '  miss'}  {c}")
    if exists and MD_FILE is None:
        MD_FILE = c

if MD_FILE is None:
    raise FileNotFoundError(
        "Could not locate GDD_Julian_Gomez.md — check the paths above "
        "and update the hardcoded candidate at the top of this script."
    )

print(f"\nSource : {MD_FILE}")
print(f"Output : {OUT_FILE}")

# ---------------------------------------------------------------------------
# Inline formatter: parses **bold**, *italic*, `code` within a string
# Returns list of (text, bold, italic, code)
# ---------------------------------------------------------------------------
INLINE_RE = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+)', re.DOTALL)

def parse_inline(text):
    runs = []
    for m in INLINE_RE.finditer(text):
        if m.group(1):
            runs.append((m.group(1), True,  False, False))
        elif m.group(2):
            runs.append((m.group(2), False, True,  False))
        elif m.group(3):
            runs.append((m.group(3), False, False, True))
        elif m.group(4):
            runs.append((m.group(4), False, False, False))
    return runs

def add_inline(para, text):
    """Add inline-formatted runs to an existing paragraph."""
    for chunk, bold, italic, code in parse_inline(text):
        run = para.add_run(chunk)
        run.bold   = bold
        run.italic = italic
        if code:
            run.font.name = "Courier New"
            run.font.size = Pt(10)

# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Base font
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# Heading styles
for lvl, size, space_before in [(1, 16, 12), (2, 13, 10), (3, 11, 8)]:
    hs = doc.styles[f"Heading {lvl}"]
    hs.font.name  = "Arial"
    hs.font.size  = Pt(size)
    hs.font.bold  = True
    hs.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    hs.paragraph_format.space_before = Pt(space_before)
    hs.paragraph_format.space_after  = Pt(4)

# ---------------------------------------------------------------------------
# Parser state
# ---------------------------------------------------------------------------
lines = MD_FILE.read_text(encoding="utf-8").splitlines()

STATE_NORMAL = "normal"
STATE_TABLE  = "table"
STATE_CODE   = "code"
STATE_LIST   = "list"

state       = STATE_NORMAL
table_rows  = []      # list of lists of cell strings
code_lines  = []
para_buf    = []      # buffered paragraph lines

def flush_para():
    """Write buffered paragraph lines as a single paragraph."""
    if not para_buf:
        return
    text = " ".join(para_buf).strip()
    if text:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_inline(p, text)
    para_buf.clear()

def flush_table():
    """Write collected table rows as a Word table."""
    if not table_rows:
        return
    cols = max(len(r) for r in table_rows)
    tbl  = doc.add_table(rows=len(table_rows), cols=cols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(table_rows):
        for ci, cell_text in enumerate(row):
            cell = tbl.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(cell_text.strip())
            run.bold = (ri == 0)   # header row bold
            run.font.size = Pt(10)
    doc.add_paragraph()   # spacing after table
    table_rows.clear()

def is_table_sep(line):
    return bool(re.match(r'\s*\|[-| :]+\|\s*$', line))

def parse_table_row(line):
    parts = line.strip().strip("|").split("|")
    return [p.strip() for p in parts]

# ---------------------------------------------------------------------------
# Main parse loop
# ---------------------------------------------------------------------------
for raw_line in lines:
    line = raw_line.rstrip()

    # --- CODE BLOCK ---
    if state == STATE_CODE:
        if line.startswith("```"):
            # End code block
            text = "\n".join(code_lines)
            if text.strip():
                p = doc.add_paragraph(text)
                p.style = "No Spacing"
                p.paragraph_format.space_after = Pt(6)
                for run in p.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
            code_lines.clear()
            state = STATE_NORMAL
        else:
            code_lines.append(line)
        continue

    if line.startswith("```"):
        flush_para()
        flush_table()
        state = STATE_CODE
        continue

    # --- TABLE ---
    if "|" in line and not line.strip().startswith("#"):
        if is_table_sep(line):
            continue  # skip separator row
        if state != STATE_TABLE:
            flush_para()
            state = STATE_TABLE
        table_rows.append(parse_table_row(line))
        continue
    else:
        if state == STATE_TABLE:
            flush_table()
            state = STATE_NORMAL

    # --- HEADINGS ---
    h4 = re.match(r'^####\s+(.*)', line)
    h3 = re.match(r'^###\s+(.*)', line)
    h2 = re.match(r'^##\s+(.*)', line)
    h1 = re.match(r'^#\s+(.*)', line)

    if h1:
        flush_para()
        doc.add_heading(h1.group(1), level=1)
        continue
    if h2:
        flush_para()
        doc.add_heading(h2.group(1), level=1)
        continue
    if h3:
        flush_para()
        doc.add_heading(h3.group(1), level=2)
        continue
    if h4:
        flush_para()
        doc.add_heading(h4.group(1), level=3)
        continue

    # --- HORIZONTAL RULE ---
    if re.match(r'^-{3,}$', line.strip()):
        flush_para()
        doc.add_paragraph("─" * 60)
        continue

    # --- BULLET LIST ---
    bm = re.match(r'^(\s*)-\s+(.*)', line)
    if bm:
        flush_para()
        indent = len(bm.group(1))
        text   = bm.group(2)
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        add_inline(p, text)
        continue

    # --- BLANK LINE ---
    if not line.strip():
        flush_para()
        continue

    # --- NORMAL PARAGRAPH ---
    para_buf.append(line)

# Flush remaining
flush_para()
flush_table()

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(str(OUT_FILE))
print(f"\n✓ Saved: {OUT_FILE.name}")
print("Open the file in Word to verify formatting.")

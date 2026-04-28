"""Idempotent: liest immer Backup, schreibt zu Original. Lücken füllen + neue Tage anhängen."""
import shutil
import unicodedata
from pathlib import Path
from docx import Document

HERE = Path(__file__).parent
SRC = HERE / "Arbeitsprotokoll_Julian_Gomez.docx"
BACKUP = HERE / "Arbeitsprotokoll_Julian_Gomez_BACKUP.docx"

if not BACKUP.exists():
    shutil.copy(SRC, BACKUP)
    print(f"Backup erstellt: {BACKUP.name}")

# IMMER vom Backup laden (idempotent)
doc = Document(BACKUP)

# --- Header-Zeitraum updaten ---
for p in doc.paragraphs:
    if "Zeitraum:" in p.text and "22.04.2026" in p.text:
        full = p.text.replace("22.04.2026", "27.04.2026")
        if p.runs:
            p.runs[0].text = full
            for r in p.runs[1:]:
                r.text = ""
        print("Header aktualisiert.")
        break

table = doc.tables[0]

# Normalisiere Strings für robusten Vergleich
def norm(s):
    return unicodedata.normalize("NFKC", s).strip()

EN_DASH = "–"
EM_DASH = "—"
EMPTY_MARKERS = {"", "-", EN_DASH, EM_DASH}

GAPS = {
    "13.03.2026": [("Pixelart: Katze idle/rise/peak — erste Sprite-Skizzen", True)],
    "18.03.2026": [("Pixelart: Stadtkulisse L1 — Gebaeude-Tiles und Fenster-Layouts", True)],
    "19.03.2026": [("Pixelart: Plattform-Tiles L1 (Beton, Geruest, Crumble-Variante)", True)],
    "20.03.2026": [("Pixelart: Hazard-Layer Smog L1 — Wolkenstruktur, Alpha-Verlauf", True)],
    "24.03.2026": [("Pixelart: Hazard-Layer Elektrizitaet L3 — Schichten, Frequenz-Varianten", True)],
    "26.03.2026": [("Pixelart: Finish-Trigger-Sprites — Windrad (L1), Glocke (L2), Hebel (L3)", True)],
    "27.03.2026": [("Pixelart: Ballon-Sammelobjekt — Frames und Schnur-Animation", True)],
    "31.03.2026": [("Pixelart: L2 Leuchtturm — Konzeptzeichnung, Silhouette, Farbpalette", True)],
    "01.04.2026": [("Pixelart: L2 Leuchtturm — Turmkoerper-Tiles und Wave-Breaker-Stufen", True)],
    "02.04.2026": [("Pixelart: L2 Wolken-Plattformen und Parallax-Layer-Skizzen", True)],
    "03.04.2026": [("Pixelart: Wespe-Sprites (Patrol-Frames, Sting, Stomp-Reaktion)", True)],
    "07.04.2026": [("krankgeschrieben (07.–16.04.2026)", False)],
    "08.04.2026": [("krankgeschrieben (07.–16.04.2026)", False)],
    "09.04.2026": [("krankgeschrieben (07.–16.04.2026)", False)],
    "10.04.2026": [("krankgeschrieben (07.–16.04.2026)", False)],
    "13.04.2026": [("krankgeschrieben (07.–16.04.2026)", False)],
    "14.04.2026": [("krankgeschrieben (07.–16.04.2026)", False)],
    "15.04.2026": [("krankgeschrieben (07.–16.04.2026)", False)],
    "16.04.2026": [("krankgeschrieben (07.–16.04.2026)", False)],
    "17.04.2026": [("Pixelart: Cat animation_sheet.png Konsolidierung (7 Frames in einem Sheet)", True)],
    "20.04.2026": [
        ("Wasp-Schwarmgroessen erhoeht: 10/15/20 pro Level (vorher 5/7/10)", True),
        ("Dialog-Bubbles: Illustrator-Plan, Bubble-Shapes definiert", True),
    ],
    "21.04.2026": [
        ("English-Migration-Audit: Texte und Identifier auf Englisch umgestellt", True),
        ("Project-Cleanup: Phase-3-Asset-Gruppen sortiert, Naming Convention durchgesetzt", True),
    ],
    "22.04.2026": [("PR #1 gemerged: Phase 04.3 L2 Elevator Interior (Collider, 404, Dialog-Code)", True)],
}

def set_cell(cell, text):
    paragraphs = cell.paragraphs
    if paragraphs:
        first = paragraphs[0]
        if first.runs:
            first.runs[0].text = text
            for r in first.runs[1:]:
                r.text = ""
        else:
            first.add_run(text)
        for p in paragraphs[1:]:
            p._element.getparent().remove(p._element)
    else:
        cell.text = text

filled = 0
for ri, row in enumerate(table.rows):
    if ri == 0:
        continue
    cells = row.cells
    date_txt = norm(cells[1].text)
    task_txt = norm(cells[2].text)
    if date_txt in GAPS and task_txt in EMPTY_MARKERS:
        entries = GAPS[date_txt]
        joined = " | ".join(e[0] for e in entries)
        is_work = entries[0][1]
        set_cell(cells[2], joined)
        if is_work:
            set_cell(cells[3], "x")
            set_cell(cells[4], "x")
            set_cell(cells[5], "x")
        filled += 1
print(f"Luecken befuellt: {filled}")

EXTRA = [
    ("Do", "23.04.2026", "Dialog-System: Damage-Title-Pools — Random pro Level mit Repeat-Guard", True),
    ("Fr", "24.04.2026", "Dialog-System: L2 Glocke + L3 Hebel — Outro-Tuning", True),
    ("Fr", "24.04.2026", "Bugfix L2: Pipe-Spalten-Alignment (pipes_top +1/-3 px)", True),
    ("Fr", "24.04.2026", "Bugfix L2: Schachtdecke-Layering (L2_ROOF_BAND_H 160→80), Dev-Browse F2", True),
    ("Fr", "24.04.2026", "Bugfix L3: Cloud-Sink Fall-Through + Crumble-Advance-Guard", True),
    ("Fr", "24.04.2026", "L3 Leuchtturm Stein-Overlay (Notfall-Fix fuer Schul-Abgabe)", True),
    ("Sa", "25.04.2026", "Push-Mechanik final verworfen (PUSH-02/03/04 dropped)", True),
    ("Sa", "25.04.2026", "Difficulty-System implementiert (Schwierigkeits-Skalierung pro Level)", True),
    ("Sa", "25.04.2026", "Z-Action umbenannt: ACTION-01 (Kletter-Kiste nur statisches Prop)", True),
    ("So", "26.04.2026", "Phase 5 Outro-Trigger: L1 Windrad, L2 Glocke, L3 Hebel — alle live", True),
    ("So", "26.04.2026", "Intro-Flow: 3-2-1 Countdown in LEVEL_INTRO, Auto-Advance ohne Klick", True),
    ("So", "26.04.2026", "audio.js v2.0: SOUNDS-Map (27 Eintraege), HTMLAudio, _onPhaseChange Hooks", True),
    ("__KW__", "", "KW 18   27.04. – 03.05.", False),
    ("Mo", "27.04.2026", "Audio-Hooks final verdrahtet (main.js _onPhaseChange Pfade)", True),
    ("Mo", "27.04.2026", "Arbeitsprotokoll vollstaendig nachgepflegt (alle Tage 04.03.–27.04.)", True),
]

added = 0
for tag, datum, aufgabe, is_work in EXTRA:
    new_row = table.add_row()
    cells = new_row.cells
    if tag == "__KW__":
        set_cell(cells[2], aufgabe)
    else:
        set_cell(cells[0], tag)
        set_cell(cells[1], datum)
        set_cell(cells[2], aufgabe)
        if is_work:
            set_cell(cells[3], "x")
            set_cell(cells[4], "x")
            set_cell(cells[5], "x")
    added += 1
print(f"Neue Zeilen angehaengt: {added}")

doc.save(SRC)
print(f"Gespeichert: {SRC.name}")

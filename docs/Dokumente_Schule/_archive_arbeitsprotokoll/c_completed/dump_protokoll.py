"""Dumpt Inhalt + Tabellenstruktur des Arbeitsprotokolls als Text."""
from docx import Document
from pathlib import Path

src = Path(__file__).parent / "Arbeitsprotokoll_Julian_Gomez.docx"
out = Path(__file__).parent / "arbeitsprotokoll_dump.txt"

doc = Document(src)
lines = []

lines.append("=" * 70)
lines.append("ABSAETZE (Body)")
lines.append("=" * 70)
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt:
        lines.append(f"[P{i:03d}] {txt}")

lines.append("")
lines.append("=" * 70)
lines.append("TABELLEN")
lines.append("=" * 70)
for ti, table in enumerate(doc.tables):
    lines.append(f"\n--- TABELLE {ti} ({len(table.rows)} Zeilen x {len(table.columns)} Spalten) ---")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip().replace("\n", " | ")
            lines.append(f"T{ti} R{ri:02d} C{ci}: {txt}")

out.write_text("\n".join(lines), encoding="utf-8")
print(f"Geschrieben: {out}")
print(f"Tabellen gefunden: {len(doc.tables)}")
print(f"Absaetze gefunden: {len(doc.paragraphs)}")
